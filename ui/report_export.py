"""
ui/report_export.py
--------------------
Turns the Writer Agent's structured output into downloadable DOCX and
PDF files. Uses python-docx and reportlab — both pure-Python, so they
work in any deployment (no LibreOffice/pandoc system dependency needed).

Requires: python-docx, reportlab (add to requirements.txt if not present).

    pip install reportlab python-docx

Expected `report` dict shape (same as the "writer" pipeline payload,
plus a "title"):
    {
        "title": str,
        "executive_summary": str,
        "key_findings": [str, ...],
        "important_concepts": [str, ...],
        "strengths": [str, ...],
        "weaknesses": [str, ...],
        "recommendations": [str, ...],
    }
"""

import re
from io import BytesIO
from xml.sax.saxutils import escape

SECTIONS = [
    ("Key Findings", "key_findings"),
    ("Important Concepts", "important_concepts"),
    ("Strengths", "strengths"),
    ("Weaknesses", "weaknesses"),
    ("Recommendations", "recommendations"),
]

# Matches markdown links: [display text](https://url)
_LINK_RE = re.compile(r"\[([^\]]+)\]\((https?://[^\s)]+)\)")


# ── Shared markdown parsing helpers ─────────────────────────────────

def _split_inline(text: str):
    """
    Splits text into ordered chunks of (text, is_bold, link_url_or_None).
    Handles **bold** and [label](url) markdown as separate inline spans,
    which covers real-world Writer/Critic output.
    """
    result = []

    pos = 0
    for m in _LINK_RE.finditer(text):
        before = text[pos:m.start()]
        if before:
            result.extend(_split_bold_only(before))
        result.append((m.group(1), False, m.group(2)))
        pos = m.end()
    remainder = text[pos:]
    if remainder:
        result.extend(_split_bold_only(remainder))

    return result


def _split_bold_only(text: str):
    """Splits text on **bold** markers, returning (chunk, is_bold, None) triples."""
    parts = re.split(r"(\*\*.*?\*\*)", text)
    result = []
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            result.append((part[2:-2], True, None))
        else:
            result.append((part, False, None))
    return result


# ── DOCX hyperlink helper ────────────────────────────────────────────

def _add_docx_hyperlink(paragraph, url: str, text: str, bold: bool = False):
    """
    python-docx has no built-in hyperlink support, so this builds the
    required run XML manually and inserts it as a real clickable,
    correctly-styled (blue + underlined) hyperlink — not just blue text.
    """
    from docx.oxml.shared import OxmlElement, qn

    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)

    if bold:
        b = OxmlElement("w:b")
        rPr.append(b)

    run.append(rPr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)

    paragraph._p.append(hyperlink)


def _add_docx_inline_runs(paragraph, text: str):
    """Adds bold runs and real clickable hyperlinks to a docx paragraph."""
    for chunk, is_bold, link_url in _split_inline(text):
        if link_url:
            _add_docx_hyperlink(paragraph, link_url, chunk, bold=is_bold)
        else:
            run = paragraph.add_run(chunk)
            run.bold = is_bold


def _add_docx_markdown(doc, text: str):
    """Parses a markdown string and appends proper Word elements
    (headings, bold runs, links, bullet/numbered lists) instead of
    dumping raw '#'/'-'/'[]()' characters into one plain paragraph."""
    if not text:
        return

    for raw_line in text.split("\n"):
        line = raw_line.strip()
        if not line:
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)", line)
        if heading_match:
            level = min(len(heading_match.group(1)) + 1, 4)  # keep below title (level 0)
            heading_text = re.sub(r"\*\*(.*?)\*\*", r"\1", heading_match.group(2)).strip()
            heading_text = _LINK_RE.sub(r"\1", heading_text)
            doc.add_heading(heading_text, level=level)
            continue

        bullet_match = re.match(r"^[-*]\s+(.*)", line)
        numbered_match = re.match(r"^\d+[\.\)]\s+(.*)", line)

        if bullet_match:
            p = doc.add_paragraph(style="List Bullet")
            item_text = bullet_match.group(1).strip()
        elif numbered_match:
            p = doc.add_paragraph(style="List Number")
            item_text = numbered_match.group(1).strip()
        else:
            p = doc.add_paragraph()
            item_text = line

        _add_docx_inline_runs(p, item_text)


# ── PDF markdown/link helpers ────────────────────────────────────────

def _to_reportlab_markup(text: str) -> str:
    """
    Converts **bold** and [label](url) markdown into reportlab's mini
    XML markup (<b>, <link href="...">), escaping everything else so
    stray '<', '>', '&' in source text don't break the PDF renderer.
    """
    out = []
    pos = 0
    for m in _LINK_RE.finditer(text):
        before = text[pos:m.start()]
        if before:
            out.append(re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", escape(before)))
        label = escape(m.group(1))
        url = escape(m.group(2))
        out.append(f'<link href="{url}" color="blue"><u>{label}</u></link>')
        pos = m.end()
    remainder = text[pos:]
    if remainder:
        out.append(re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", escape(remainder)))
    return "".join(out)


def _markdown_to_pdf_flowables(text: str, styles):
    """Parses markdown into a list of reportlab flowables (headings,
    bullet/numbered lists, paragraphs) instead of one raw text blob."""
    from reportlab.platypus import ListFlowable, ListItem, Paragraph, Spacer

    if not text:
        return []

    flowables = []
    bullet_buffer = []
    number_buffer = []

    def flush_bullets():
        if bullet_buffer:
            items = [ListItem(Paragraph(t, styles["Normal"])) for t in bullet_buffer]
            flowables.append(ListFlowable(items, bulletType="bullet"))
            flowables.append(Spacer(1, 8))
            bullet_buffer.clear()

    def flush_numbers():
        if number_buffer:
            items = [ListItem(Paragraph(t, styles["Normal"])) for t in number_buffer]
            flowables.append(ListFlowable(items, bulletType="1"))
            flowables.append(Spacer(1, 8))
            number_buffer.clear()

    heading_styles = {1: "Heading1", 2: "Heading2", 3: "Heading3", 4: "Heading4"}

    for raw_line in text.split("\n"):
        line = raw_line.strip()
        if not line:
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)", line)
        if heading_match:
            flush_bullets()
            flush_numbers()
            level = min(len(heading_match.group(1)), 4)
            heading_text = re.sub(r"\*\*(.*?)\*\*", r"\1", heading_match.group(2)).strip()
            heading_text = _LINK_RE.sub(r"\1", heading_text)
            flowables.append(Paragraph(escape(heading_text), styles[heading_styles[level]]))
            flowables.append(Spacer(1, 6))
            continue

        bullet_match = re.match(r"^[-*]\s+(.*)", line)
        numbered_match = re.match(r"^\d+[\.\)]\s+(.*)", line)

        if bullet_match:
            flush_numbers()
            bullet_buffer.append(_to_reportlab_markup(bullet_match.group(1).strip()))
            continue
        if numbered_match:
            flush_bullets()
            number_buffer.append(_to_reportlab_markup(numbered_match.group(1).strip()))
            continue

        flush_bullets()
        flush_numbers()
        flowables.append(Paragraph(_to_reportlab_markup(line), styles["Normal"]))
        flowables.append(Spacer(1, 4))

    flush_bullets()
    flush_numbers()
    return flowables


# ── DOCX ─────────────────────────────────────────────────────────────

def generate_docx(report: dict) -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading(report.get("title", "Research Report"), level=0)

    doc.add_heading("Executive Summary", level=1)
    _add_docx_markdown(doc, report.get("executive_summary", ""))

    for heading, key in SECTIONS:
        items = report.get(key, [])
        if not items:
            continue
        doc.add_heading(heading, level=1)
        for item in items:
            p = doc.add_paragraph(style="List Bullet")
            _add_docx_inline_runs(p, item)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── PDF ──────────────────────────────────────────────────────────────

def generate_pdf(report: dict) -> bytes:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer

    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter, title=report.get("title", "Research Report"))
    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph(escape(report.get("title", "Research Report")), styles["Title"]))
    story.append(Spacer(1, 16))

    story.append(Paragraph("Executive Summary", styles["Heading1"]))
    story.extend(_markdown_to_pdf_flowables(report.get("executive_summary", ""), styles))
    story.append(Spacer(1, 12))

    for heading, key in SECTIONS:
        items = report.get(key, [])
        if not items:
            continue
        story.append(Paragraph(heading, styles["Heading1"]))
        bullets = [ListItem(Paragraph(_to_reportlab_markup(item), styles["Normal"])) for item in items]
        story.append(ListFlowable(bullets, bulletType="bullet"))
        story.append(Spacer(1, 12))

    doc.build(story)
    return buf.getvalue()